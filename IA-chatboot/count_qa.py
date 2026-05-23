import docx
import re

doc = docx.Document('ADI.docx')
num_paras = len(doc.paragraphs)
print(f'Total paragraphs: {num_paras}')

# Save all non-empty paragraphs
with open('all_paras.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            f.write(f'{i}: {txt}\n')

print('Saved all paragraphs to all_paras.txt')

# Read back and search
with open('all_paras.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Find all "question_en" fields
questions = re.findall(r'"question_en":\s*"([^"]+)"', content)
print(f'\nTotal question_en fields: {len(questions)}')
if questions:
    print('All questions:')
    for i, q in enumerate(questions, 1):
        print(f'{i}. {q}')

# Find corresponding answers
answers = re.findall(r'"answer_en":\s*"([^"]+)"', content)
print(f'\nTotal answer_en fields: {len(answers)}')
if answers:
    print('All answers:')
    for i, a in enumerate(answers, 1):
        print(f'{i}. {a}')
