import docx

doc = docx.Document('ADI.docx')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print('\n--- Sample raw paragraphs (first 20) ---')
for i, p in enumerate(doc.paragraphs[:20]):
    if p.text.strip():
        print(f'{i}: {p.text[:100]}')

# Check for actual question/answer structure
text = '\n'.join([p.text for p in doc.paragraphs])
print(f'\n--- Searching for Q&A patterns ---')
lines = text.split('\n')
questions = [l for l in lines if '?' in l and len(l) < 200]
print(f'Lines with "?": {len(questions)}')
if questions:
    print('First 3 question-like lines:')
    for q in questions[:3]:
        print(f'  - {q[:80]}')
