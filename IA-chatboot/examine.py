import docx

doc = docx.Document('ADI.docx')
print(f'Paragraphs: {len(doc.paragraphs)}')

# Print raw text content of first 100 paragraphs that have content
count = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt:
        count += 1
        if count <= 100:
            # Write to file to avoid encoding issues
            with open('paras.txt', 'a', encoding='utf-8') as f:
                f.write(f'{count}: {txt}\n')

print(f'Non-empty paragraphs: {count}')
print('First 30 saved to paras.txt')

# Also check runs to see if JSON is in a specific paragraph
print('\nSearching for question patterns...')
with open('paras.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Look for question-like entries
import re
questions = re.findall(r'"question[^"]*":\s*"([^"]{10,200})"', content)
print(f'Found {len(questions)} question strings')
if questions:
    print('Sample questions:')
    for i, q in enumerate(questions[:10]):
        print(f'{i+1}. {q}')

answers = re.findall(r'"answer[^"]*":\s*"([^"]{10,200})"', content)
print(f'\nFound {len(answers)} answer strings')
